from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 设置默认字体 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 辅助函数 =====
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, size=None, indent=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (indent_level + 1))
    return p

def set_cell_font(cell, text, bold=False, size=Pt(10)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def shade_cells(row, color_hex):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color_hex,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

# ===== 标题 =====
title = doc.add_heading('聊天软件增强版 - 更改清单', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()  # 空行

# ===== 一、新增功能（27项） =====
add_heading_styled('一、新增功能（27项）', level=1)

# 消息类型扩展
add_heading_styled('消息类型扩展', level=2)
table1 = doc.add_table(rows=5, cols=3, style='Table Grid')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers1 = ['序号', '功能', '说明']
for i, h in enumerate(headers1):
    set_cell_font(table1.rows[0].cells[i], h, bold=True)
shade_cells(table1.rows[0], '4472C4')
for cell in table1.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data1 = [
    ['1', '表情消息', '96个表情选择面板'],
    ['2', '图片消息', '6张示例图片'],
    ['3', '语音消息', '录音计时发送'],
    ['4', '系统消息', '撤回提示等'],
]
for r_idx, row_data in enumerate(data1):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table1.rows[r_idx + 1].cells[c_idx], val)

# 消息操作功能
add_heading_styled('消息操作功能', level=2)
table2 = doc.add_table(rows=7, cols=3, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table2.rows[0].cells[i], h, bold=True)
shade_cells(table2.rows[0], '4472C4')
for cell in table2.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data2 = [
    ['5', '消息撤回', '2分钟内可撤回'],
    ['6', '消息复制', '复制到剪贴板'],
    ['7', '消息转发', '转发给其他联系人'],
    ['8', '消息引用', '引用回复消息'],
    ['9', '消息删除', '删除单条消息'],
    ['10', '长按菜单', '显示操作菜单'],
]
for r_idx, row_data in enumerate(data2):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table2.rows[r_idx + 1].cells[c_idx], val)

# 消息状态功能
add_heading_styled('消息状态功能', level=2)
table3 = doc.add_table(rows=5, cols=3, style='Table Grid')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table3.rows[0].cells[i], h, bold=True)
shade_cells(table3.rows[0], '4472C4')
for cell in table3.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data3 = [
    ['11', '已发送状态', '显示"已发送"'],
    ['12', '已读状态', '显示"已读"'],
    ['13', '未读计数', '显示未读消息数量'],
    ['14', '标记未读', '手动标记为未读'],
]
for r_idx, row_data in enumerate(data3):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table3.rows[r_idx + 1].cells[c_idx], val)

# 时间显示功能
add_heading_styled('时间显示功能', level=2)
table4 = doc.add_table(rows=4, cols=3, style='Table Grid')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table4.rows[0].cells[i], h, bold=True)
shade_cells(table4.rows[0], '4472C4')
for cell in table4.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data4 = [
    ['15', '消息时间戳', '每条消息显示时间'],
    ['16', '时间格式化', 'HH:MM格式显示'],
    ['17', '完整时间', 'MM-DD HH:MM格式'],
]
for r_idx, row_data in enumerate(data4):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table4.rows[r_idx + 1].cells[c_idx], val)

# 搜索功能
add_heading_styled('搜索功能', level=2)
table5 = doc.add_table(rows=4, cols=3, style='Table Grid')
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table5.rows[0].cells[i], h, bold=True)
shade_cells(table5.rows[0], '4472C4')
for cell in table5.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data5 = [
    ['18', '消息搜索', '搜索历史消息'],
    ['19', '搜索面板', '显示搜索结果'],
    ['20', '实时搜索', '输入即搜索'],
]
for r_idx, row_data in enumerate(data5):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table5.rows[r_idx + 1].cells[c_idx], val)

# 数据保存功能
add_heading_styled('数据保存功能', level=2)
table6 = doc.add_table(rows=3, cols=3, style='Table Grid')
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table6.rows[0].cells[i], h, bold=True)
shade_cells(table6.rows[0], '4472C4')
for cell in table6.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data6 = [
    ['21', '草稿保存', '自动保存输入内容'],
    ['22', '草稿恢复', '返回时恢复草稿'],
]
for r_idx, row_data in enumerate(data6):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table6.rows[r_idx + 1].cells[c_idx], val)

# 交互提示功能
add_heading_styled('交互提示功能', level=2)
table7 = doc.add_table(rows=3, cols=3, style='Table Grid')
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table7.rows[0].cells[i], h, bold=True)
shade_cells(table7.rows[0], '4472C4')
for cell in table7.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data7 = [
    ['23', '正在输入提示', '显示"正在输入..."'],
    ['24', '输入状态模拟', '模拟对方输入状态'],
]
for r_idx, row_data in enumerate(data7):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table7.rows[r_idx + 1].cells[c_idx], val)

# 更多功能
add_heading_styled('更多功能', level=2)
table8 = doc.add_table(rows=4, cols=3, style='Table Grid')
table8.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers1):
    set_cell_font(table8.rows[0].cells[i], h, bold=True)
shade_cells(table8.rows[0], '4472C4')
for cell in table8.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

data8 = [
    ['25', '清空聊天记录', '一键清空所有消息'],
    ['26', '更多功能面板', '图片、搜索等入口'],
    ['27', '右上角菜单', '搜索、标记未读、清空'],
]
for r_idx, row_data in enumerate(data8):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table8.rows[r_idx + 1].cells[c_idx], val)

# ===== 二、数据模型扩展 =====
add_heading_styled('二、数据模型扩展', level=1)

add_heading_styled('消息对象新增字段', level=2)
table_msg = doc.add_table(rows=9, cols=2, style='Table Grid')
table_msg.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_msg.rows[0].cells[0], '字段', bold=True)
set_cell_font(table_msg.rows[0].cells[1], '说明', bold=True)
shade_cells(table_msg.rows[0], '4472C4')
for cell in table_msg.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

msg_fields = [
    ['id', '消息唯一ID'],
    ['type', '消息类型（文本/表情/图片/语音/系统）'],
    ['duration', '语音时长'],
    ['timestamp', '时间戳'],
    ['status', '消息状态（发送中/已发送/已送达/已读/失败）'],
    ['isRecalled', '是否已撤回'],
    ['isRead', '是否已读'],
    ['replyTo', '引用回复的消息'],
]
for r_idx, row_data in enumerate(msg_fields):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_msg.rows[r_idx + 1].cells[c_idx], val)

add_heading_styled('联系人对象新增字段', level=2)
table_item = doc.add_table(rows=7, cols=2, style='Table Grid')
table_item.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_item.rows[0].cells[0], '字段', bold=True)
set_cell_font(table_item.rows[0].cells[1], '说明', bold=True)
shade_cells(table_item.rows[0], '4472C4')
for cell in table_item.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

item_fields = [
    ['unreadCount', '未读消息数'],
    ['lastMessageTime', '最后消息时间'],
    ['isTyping', '对方是否正在输入'],
    ['draft', '草稿内容'],
    ['isGroup', '是否为群聊'],
    ['members', '群成员列表'],
]
for r_idx, row_data in enumerate(item_fields):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_item.rows[r_idx + 1].cells[c_idx], val)

# ===== 三、功能问题修复（2项） =====
add_heading_styled('三、功能问题修复（2项）', level=1)
table_fix = doc.add_table(rows=3, cols=2, style='Table Grid')
table_fix.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_fix.rows[0].cells[0], '序号', bold=True)
set_cell_font(table_fix.rows[0].cells[1], '说明', bold=True)
shade_cells(table_fix.rows[0], '4472C4')
for cell in table_fix.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

fix_data = [
    ['1', '录音功能修复 - 添加停止录音按钮'],
    ['2', '草稿保存修复 - 使用静态Map持久化存储'],
]
for r_idx, row_data in enumerate(fix_data):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_fix.rows[r_idx + 1].cells[c_idx], val)

# ===== 四、编译错误修复（31个） =====
add_heading_styled('四、编译错误修复（31个）', level=1)
table_compile = doc.add_table(rows=6, cols=2, style='Table Grid')
table_compile.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_compile.rows[0].cells[0], '批次', bold=True)
set_cell_font(table_compile.rows[0].cells[1], '说明', bold=True)
shade_cells(table_compile.rows[0], '4472C4')
for cell in table_compile.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

compile_data = [
    ['第1批', 'null类型错误（8个）'],
    ['第2批', 'timestamp undefined错误（3个）'],
    ['第3批', 'Item缺少字段错误（13个）'],
    ['第4批', 'Row space属性错误（1个）'],
    ['第5批', '函数参数null错误（6个）'],
]
for r_idx, row_data in enumerate(compile_data):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_compile.rows[r_idx + 1].cells[c_idx], val)

# ===== 五、UI优化 =====
add_heading_styled('五、UI优化', level=1)

add_heading_styled('消息气泡样式', level=2)
table_ui = doc.add_table(rows=5, cols=3, style='Table Grid')
table_ui.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_ui.rows[0].cells[0], '类型', bold=True)
set_cell_font(table_ui.rows[0].cells[1], '对方', bold=True)
set_cell_font(table_ui.rows[0].cells[2], '自己', bold=True)
shade_cells(table_ui.rows[0], '4472C4')
for cell in table_ui.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

ui_data = [
    ['文本', '灰色背景', '绿色背景'],
    ['表情', '32px大号字体', '32px大号字体'],
    ['图片', '100x100尺寸', '100x100尺寸'],
    ['语音', '麦克风图标+时长', '时长+麦克风图标'],
]
for r_idx, row_data in enumerate(ui_data):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_ui.rows[r_idx + 1].cells[c_idx], val)

add_heading_styled('新增UI组件', level=2)
ui_components = [
    '表情选择面板（8列x12行）',
    '图片选择面板（3列x2行）',
    '更多功能面板',
    '录音提示遮罩',
    '上下文菜单',
    '搜索面板',
    '转发面板',
]
for comp in ui_components:
    add_bullet(comp)

# ===== 六、项目统计 =====
add_heading_styled('六、项目统计', level=1)
table_stat = doc.add_table(rows=6, cols=2, style='Table Grid')
table_stat.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_stat.rows[0].cells[0], '项目', bold=True)
set_cell_font(table_stat.rows[0].cells[1], '数值', bold=True)
shade_cells(table_stat.rows[0], '4472C4')
for cell in table_stat.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

stat_data = [
    ['新增功能', '27项'],
    ['数据模型字段', '16个'],
    ['编译错误修复', '31个'],
    ['功能问题修复', '2个'],
    ['代码行数', '800+行'],
]
for r_idx, row_data in enumerate(stat_data):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_stat.rows[r_idx + 1].cells[c_idx], val)

# ===== 七、版本信息 =====
add_heading_styled('七、版本信息', level=1)
table_ver = doc.add_table(rows=6, cols=2, style='Table Grid')
table_ver.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table_ver.rows[0].cells[0], '项目', bold=True)
set_cell_font(table_ver.rows[0].cells[1], '信息', bold=True)
shade_cells(table_ver.rows[0], '4472C4')
for cell in table_ver.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

ver_data = [
    ['原版本', 'simple-chat-list-master'],
    ['增强版本', 'simple-chat-list-enhanced v2.0'],
    ['更新日期', '2026-05-12'],
    ['HarmonyOS版本', '5.0.5 Release'],
    ['项目状态', '编译成功，功能完整'],
]
for r_idx, row_data in enumerate(ver_data):
    for c_idx, val in enumerate(row_data):
        set_cell_font(table_ver.rows[r_idx + 1].cells[c_idx], val)

# ===== 保存文件 =====
output_path = r'd:\CODE\simple-chat-list-enhanced\更改清单_精简版.docx'
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
